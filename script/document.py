from docx import Document

def create_word_doc(book_title, content_list, detailed_info_list):
    doc = Document()
    doc.add_heading('Table of Content', level=0)
    for item in content_list:
        doc.add_paragraph(item)

    for item, info in zip(content_list, detailed_info_list):
        doc.add_heading(item)
        doc.add_paragraph(info)

    doc.save(f'{book_title}.docx')
